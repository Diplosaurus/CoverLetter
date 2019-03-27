#!/usr/bin/env python

import os
import sys
import shutil
import datetime

from docx import Document, shared


##############################
# Gets the current directory #
##############################
current_path = os.getcwd()

########################
# Initialize variables #
########################
COVER_LETTER = "Cover Letter.docx"
recipient = sys.argv[1]
RECIPIENT_IDENTIFIER = "[RECIPIENT]"
DATE_IDENTIFIER = "[DATE]"
DESTINATION_FOLDER_NAME = "Cover_Letters"

################################################
# Creates a document object from the docx file #
################################################
document = Document(COVER_LETTER)


################################################
# Creates a new folder if it doesn't exist yet #
################################################
def create_folder(folder_name):
    new_path = current_path + "/" + folder_name
    if not os.path.exists(new_path):
            os.makedirs(folder_name)

###########################################################
# Formats the paragraph at INDEX using FONT_NAME and SIZE #
###########################################################
def formatParagraph(index, font_name, size):
    to_format = document.paragraphs[index].runs[0]
    to_format.font.name = font_name
    to_format.font.size = shared.Pt(size)

#######################################
# Gets today's date in Mon. Day, Year #
#######################################
def get_date():
    date = datetime.date.today().ctime().split(" ")
    return f"{date[1]}. {date[2]}, {date[len(date) - 1]}"


#################################################
# Traverses docx file and edits necessary parts #
#################################################
for index in range(len(document.paragraphs)):
    line = document.paragraphs[index].text
    if RECIPIENT_IDENTIFIER in line:#magic_word1 in line and magic_word2 in line:
        document.paragraphs[index].text = document.paragraphs[index].text.replace(RECIPIENT_IDENTIFIER, recipient)#magic_word1 + " " + recipient + " " + magic_word2
        formatParagraph(index, "Times new Roman", 12)
    
    elif DATE_IDENTIFIER in line:
        document.paragraphs[index].text = get_date()
        formatParagraph(index, "Times new Roman", 12)


####################################################
# Saves the new docx file to the destintion folder #
####################################################
create_folder(DESTINATION_FOLDER_NAME)
new_doc_name = recipient + " Cover Letter.docx"
document.save(new_doc_name)

if not os.path.exists(current_path + "/" + DESTINATION_FOLDER_NAME + "/" + new_doc_name):
    shutil.move(current_path + "/" + new_doc_name, current_path + "/" + DESTINATION_FOLDER_NAME)
else:
    os.remove(current_path + "/" + DESTINATION_FOLDER_NAME +  "/" + new_doc_name)
    shutil.move(current_path + "/" + new_doc_name, current_path + "/" + DESTINATION_FOLDER_NAME)






    
